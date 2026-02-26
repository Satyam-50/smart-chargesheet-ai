import os
import io
import tempfile
from pathlib import Path
import importlib.util

try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_pdf(file_bytes):
    """Extract text from PDF using pypdf"""
    if not PYPDF_AVAILABLE:
        return {"error": "PyPDF not installed. Run: pip install pypdf"}
    
    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += f"\n--- PAGE {page_num + 1} ---\n"
            text += page.extract_text()
        return {"text": text, "pages": len(pdf_reader.pages)}
    except Exception as e:
        return {"error": f"PDF extraction failed: {str(e)}"}


def extract_text_from_docx(file_bytes):
    """Extract text from DOCX"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        return {"text": text, "paragraphs": len(doc.paragraphs)}
    except Exception as e:
        return {"error": f"DOCX extraction failed: {str(e)}"}


def extract_text_from_image_ocr(file_bytes, lang='hin+eng'):
    """Extract text from image using Tesseract OCR"""
    if not PYTESSERACT_AVAILABLE:
        return {"error": "pytesseract/Pillow not installed. Run: pip install pytesseract Pillow"}
    
    # Check if Tesseract is installed on system
    try:
        pytesseract.get_tesseract_version()
    except Exception:
        return {"error": "Tesseract OCR not installed on system. Install from: https://github.com/UB-Mannheim/tesseract/wiki"}
    
    try:
        image = Image.open(io.BytesIO(file_bytes))
        # lang parameter: 'hin' for Hindi, 'eng' for English, 'hin+eng' for both
        text = pytesseract.image_to_string(image, lang=lang)
        return {"text": text, "language": lang}
    except Exception as e:
        return {"error": f"OCR extraction failed: {str(e)}"}


def extract_text_from_txt(file_bytes):
    """Extract text from plain text file"""
    try:
        text = file_bytes.decode('utf-8', errors='ignore')
        return {"text": text, "encoding": "utf-8"}
    except Exception as e:
        return {"error": f"Text extraction failed: {str(e)}"}


def extract_text_from_file(file_bytes, filename):
    """
    Universal file handler - routes to appropriate extractor based on file type
    Returns: { "text": extracted_text, "format": file_format, ... }
    """
    file_ext = Path(filename).suffix.lower()
    
    if file_ext == ".pdf":
        result = extract_text_from_pdf(file_bytes)
        result["format"] = "pdf"
        return result
    
    elif file_ext == ".docx":
        result = extract_text_from_docx(file_bytes)
        result["format"] = "docx"
        return result
    
    elif file_ext in [".txt", ".log"]:
        result = extract_text_from_txt(file_bytes)
        result["format"] = "txt"
        return result
    
    elif file_ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif"]:
        result = extract_text_from_image_ocr(file_bytes)
        result["format"] = "image"
        return result
    
    else:
        return {"error": f"Unsupported file format: {file_ext}. Supported: PDF, DOCX, TXT, PNG, JPG, JPEG, BMP, GIF"}
